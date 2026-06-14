# -*- coding: utf-8 -*-
"""HID 과제 산출물 생성: 1 Pager + CoT Note (질문 북마크 타임라인)"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

KFONT = "Malgun Gothic"

def base_doc():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = KFONT
    st.font.size = Pt(10.5)
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), KFONT)
    # heading fonts -> Korean
    for h in ["Heading 1", "Heading 2", "Heading 3", "Title"]:
        try:
            s = doc.styles[h]
            s.font.name = KFONT
            r = s.element.get_or_add_rPr()
            rf = r.get_or_add_rFonts()
            rf.set(qn("w:eastAsia"), KFONT)
        except KeyError:
            pass
    return doc

def set_run_kfont(run):
    run.font.name = KFONT
    rpr = run._element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn("w:eastAsia"), KFONT)

def p(doc, text="", bold=False, size=10.5, color=None, space_after=6, align=None):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space_after)
    if align:
        par.alignment = align
    run = par.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    set_run_kfont(run)
    return par

def bullet(doc, text, bold_lead=None):
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r1 = par.add_run(bold_lead)
        r1.bold = True
        set_run_kfont(r1)
        r2 = par.add_run(text)
        set_run_kfont(r2)
    else:
        r = par.add_run(text)
        set_run_kfont(r)
    return par

def section_bar(doc, text):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(10)
    par.paragraph_format.space_after = Pt(6)
    run = par.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    set_run_kfont(run)

def numbered_head(doc, text):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(8)
    par.paragraph_format.space_after = Pt(3)
    run = par.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    set_run_kfont(run)

def make_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        set_run_kfont(run)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9.5)
            set_run_kfont(run)
    return t

# ============================================================
# 1) 1 PAGER
# ============================================================
doc = base_doc()

title = p(doc, "질문 북마크 타임라인 (Question Bookmark Timeline)", bold=True, size=18, space_after=2)
p(doc, "HID 1 Pager · v0.1 · 듣는 흐름을 끊지 않고 ‘막힌 순간’을 질문으로 저장하는 LLM 채팅 인터랙션",
  size=10, color=(0x66, 0x66, 0x66), space_after=4)

section_bar(doc, "[1 Pager] 문제 정의")

numbered_head(doc, "1. 사용자와 페인포인트")
p(doc, "온라인 강의·실시간 회의·웨비나를 들으며 LLM을 보조 학습 도구로 함께 쓰는 주니어 실무자·학습자. "
       "재생 속도를 스스로 조절하기 어려운 ‘실시간으로 흐르는’ 환경에서, 매번 LLM을 옆에 띄워두고 듣는다. "
       "이해가 막히는 순간이 생기면 ‘지금 질문하면 다음 내용을 놓치고, 나중에 질문하면 방금의 맥락을 잊는’ "
       "딜레마를 반복적으로 겪는다. 정작 본인은 이 충돌을 ‘내 집중력 문제’ 또는 ‘어쩔 수 없는 일’로 여겨, "
       "마찰로 인식하지 못한 상태다.")

numbered_head(doc, "2. 사용자의 목표")
p(doc, "실시간으로 흐르는 강의/회의의 흐름을 놓치지 않으면서, 이해가 막힌 지점을 ‘그 순간의 맥락과 함께’ "
       "표시해 두고, 끝난 뒤(또는 잠깐 멈춘 틈에) 그 맥락을 기반으로 질문해 이해를 메우고 복습한다. "
       "완료 조건은 관찰 가능하다: 막힌 지점 N개를 흐름을 끊지 않고 표시 → 각 지점을 질문/답변 카드로 해결 "
       "→ 타임라인으로 ‘내가 어디서 막혔는지’를 시간순으로 복기.")

numbered_head(doc, "3. 핵심 마찰 정의")
p(doc, "채팅 UX는 질문이 발생한 ‘시간 맥락’을 보존하지 못한다. 질문이 성립하려면 ‘방금 그 부분’이라는 맥락이 "
       "필요한데, 선형 채팅에서는 그 맥락을 사용자가 직접 문장으로 다시 번역해 입력해야 한다. 그 결과 사용자는 "
       "(a) 지금 질문을 작성하느라 현재 흐름을 놓치거나, (b) 나중에 질문하려고 방금의 맥락을 기억해 재구성해야 하는, "
       "‘듣기’와 ‘질문 만들기’가 충돌하는 상황에 놓인다. 이는 모델 답변 품질의 문제가 아니라, 질문 입력이 "
       "‘시간 위의 한 지점’이 아니라 ‘맥락 없는 독립 텍스트’로만 다뤄지는 인터랙션 구조의 문제다. "
       "(음성 입력으로 바꿔도 ‘여기/아까 그 부분’이 무엇을 가리키는지는 그대로 미해결로 남는다.)")

section_bar(doc, "[1 Pager] 해결책")

numbered_head(doc, "4. 인터랙션 설계")
p(doc, "대화창과 분리된 ‘질문 타임라인’을 두고, 질문 입력을 ‘문장 작성’이 아니라 ‘시간 위의 한 지점 표시’로 바꾼다. "
       "사용자의 행위와 시스템의 반응을 쌍으로 설계한다.", space_after=4)
bullet(doc, "강의/회의가 실시간으로 흐르는 중 이해가 막히면 → 사용자가 ‘막힘’ 버튼을 한 번 탭(또는 Q키)한다.")
bullet(doc, "탭하는 순간 → 시스템은 흐름을 멈추지 않고, 직전 20~30초 transcript를 snapshot으로 캡처해 "
            "타임라인 위에 ‘막힌 지점 카드’를 핀으로 고정한다. (지금 질문을 만들 필요가 없으므로 흐름이 유지된다.)")
bullet(doc, "사용자가 나중에 카드를 열면 → 캡처된 맥락 요약과, 그 맥락에 맞춘 질문 후보 3개가 제시된다.")
bullet(doc, "사용자가 후보를 선택하거나 한 줄로 다듬으면 → AI 답변이 채팅 로그가 아니라 해당 타임스탬프 카드 아래에 고정된다.")
bullet(doc, "사용자가 복습 모드를 열면 → 막힌 지점들이 시간순으로 정렬되고, 해결/미해결 상태와 Q&A가 함께 복기된다.")
p(doc, "", space_after=2)
p(doc, "상태 전이 (입력/조건 → 시스템 출력)", bold=True, size=10.5, space_after=4)
make_table(doc,
    ["상태", "사용자 입력 / 조건", "시스템 출력"],
    [
        ["Listening", "강의/회의를 듣는 중", "현재 transcript가 시간순으로 흐름"],
        ["Marked", "‘막힘’ 버튼 탭(Q)", "직전 20~30초 맥락을 snapshot으로 저장, 타임라인에 핀 고정"],
        ["Question Candidate", "북마크 카드 열람", "맥락 요약 + 맥락에 맞춘 질문 후보 3개 표시"],
        ["Answered", "질문 후보 선택/수정", "답변 카드 생성, 해당 타임스탬프에 고정"],
        ["Review", "나중에 타임라인 재방문", "막힌 지점과 답변을 시간순으로 복습(해결/미해결 표시)"],
    ])

numbered_head(doc, "5. 테스트 시나리오")
p(doc, "프로토타입을 처음 보는 사람이 따라가면, 기존 채팅 방식과의 before/after를 직접 체감할 수 있다.", space_after=4)
steps = [
    "더미 강의(‘A/B 테스트 기초: 통계적 유의성’) transcript가 재생되는 것을 확인한다.",
    "(대조군) 상단에서 ‘기존 채팅 방식’으로 전환해 질문을 입력해 본다 → 입력하는 동안에도 자막이 계속 흐르고, "
    "‘질문 작성 중 놓친 자막 N줄’이 집계되는 것을 확인한다. (마찰 체감)",
    "‘질문 북마크 방식’으로 전환한다.",
    "강의가 흐르는 중 이해가 막히는 순간 ‘막힘’ 버튼(또는 Q)을 누른다 → 흐름은 멈추지 않고, "
    "타임라인에 직전 맥락이 담긴 카드가 생기는 것을 확인한다.",
    "막힌 지점을 2~3개 더 표시한다.",
    "카드를 열어 자동 저장된 맥락 요약과 질문 후보 3개를 확인하고, 하나를 선택하거나 직접 다듬어 보낸다 "
    "→ 답변이 해당 타임스탬프 카드에 붙는 것을 확인한다.",
    "‘복습 모드’를 열어 막힌 지점·해결 여부·Q&A가 시간순으로 정리되는 것을 확인한다.",
]
for i, s in enumerate(steps, 1):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(3)
    r = par.add_run(f"{i}. {s}")
    r.font.size = Pt(10.5)
    set_run_kfont(r)

numbered_head(doc, "6. 프로토타입 링크")
bullet(doc, "https://question-bookmark-timeline.vercel.app — Next.js(App Router)로 구현해 "
            "Vercel에 배포한 인터랙티브 프로토타입(공개 접속 가능).",
       bold_lead="Vercel URL: ")
bullet(doc, "question-bookmark-timeline/ (Next.js 프로젝트). 로컬 실행은 npm install 후 npm run dev "
            "(http://localhost:3000). 단일 HTML 버전은 질문북마크_타임라인_과제/prototype.html.",
       bold_lead="소스: ")
bullet(doc, "STT 없이 더미 transcript로 핵심을 재현(v0.1). ‘기존 채팅 / 질문 북마크’ 모드 토글로 "
            "before/after를 한 화면에서 비교할 수 있다.", bold_lead="범위: ")
bullet(doc, "‘실시간 흐름을 끊지 않고 질문이 발생한 시간 맥락을 저장·복습하는가’를 증명하는 것.", bold_lead="검증 포인트: ")

doc.save("1Pager_질문북마크타임라인.docx")
print("saved 1Pager")

# ============================================================
# 2) CoT NOTE
# ============================================================
doc = base_doc()

p(doc, "CoT Note — 질문 북마크 타임라인", bold=True, size=18, space_after=2)
p(doc, "Abduction → AI → Gap Analysis → Takeaway", size=10, color=(0x66, 0x66, 0x66), space_after=6)

section_bar(doc, "1. Abduction: 나의 답변 먼저 써보기")
p(doc, "[나의 첫 문제 정의]", bold=True, space_after=3)
p(doc, "“LLM에게 질문할 때 긴 텍스트를 타이핑하는 게 번거롭다. 음성으로 질문하게 하면 마찰이 줄지 않을까?” "
       "— 마찰의 원인을 ‘입력이 귀찮다’로 보고, 입력 채널을 텍스트에서 음성으로 바꾸는 방향을 떠올렸다.")
p(doc, "[나의 첫 해결책 스케치]", bold=True, space_after=3)
bullet(doc, "채팅 입력창에 음성 버튼을 추가하고 STT로 받아쓴다.")
bullet(doc, "강의를 듣다 모르는 게 생기면 음성으로 바로 물어본다.")
p(doc, "[막힌 점]", bold=True, space_after=3)
p(doc, "‘음성으로 바꾸면 편해진다’는 직관은 있었지만, “왜 텍스트가 불편한가”를 인터랙션 구조 차원에서 "
       "설명하지 못했다. 마찰을 ‘입력 노동’으로만 봤기 때문에, 답이 늘 ‘입력을 더 쉽게’로 수렴했다. "
       "강의를 들으며 LLM을 쓴 경험은 있었지만, 그 안의 충돌 장면을 마찰의 재료로 꺼내지 못했다.")

section_bar(doc, "2. AI: AI에게 답변 받기")
p(doc, "[요청 방식]", bold=True, space_after=3)
p(doc, "정답만 달라고 하지 않고, 내 시도(‘음성 입력’)를 보여주며 “이건 입력 채널만 바꾼 것 같다. "
       "단순 음성 입력이 아니라 더 고맥락의 문제로 파고든 후보들을 달라”고 요청했다.")
p(doc, "[AI가 드러낸 것]", bold=True, space_after=3)
bullet(doc, "음성으로 말해도 ‘여기/아까 그 부분/이 흐름’ 같은 지시어가 무엇을 가리키는지는 여전히 구조화되지 않는다.")
bullet(doc, "따라서 진짜 마찰은 ‘말하기 어려움’이 아니라 ‘맥락을 입력 구조로 만들지 못하는 것’이다.")
bullet(doc, "이 관점으로 후보 10개를 펼쳤고(질문 북마크 타임라인, Point & Ask, Voice Dump→Structure Canvas 등), "
            "공통 축은 ‘맥락을 입력으로 가져오기’였다. 적합도·프로토타입 가능성·확장성 기준에서 "
            "‘질문 북마크 타임라인’이 1순위로 도출됐다.")
p(doc, "[내 시도에 대한 메타 진단]", bold=True, space_after=3)
bullet(doc, "내 시도를 ‘음성 입력’이라고 구체적으로 적어 보여준 것. 덕분에 AI가 "
            "“그건 채널 변경에 그친다”고 정확히 짚을 수 있었다.", bold_lead="잘된 점: ")
bullet(doc, "사용 장면을 한 컷으로 좁히지 못했다. ‘강의를 듣다가’까지는 갔지만, "
            "‘질문을 만드는 순간 현재 흐름을 놓친다’는 충돌 장면을 못 봤다.", bold_lead="빠진 것: ")
bullet(doc, "‘마찰 = 입력의 번거로움’으로 좁게 정의했기 때문. 마찰을 입력 노동으로만 보면 "
            "해결책은 항상 ‘입력을 더 쉽게’로 귀결된다.", bold_lead="왜 그랬는가: ")

section_bar(doc, "3. Gap Analysis: 내 답변과 비교 분석해보기")
p(doc, "Gap 1. 마찰의 위치", bold=True, space_after=3)
bullet(doc, "마찰은 ‘입력 방식(텍스트)’에 있다 → 음성으로 바꾸면 해결된다.", bold_lead="나의 시도: ")
bullet(doc, "마찰은 ‘입력 채널’이 아니라 ‘맥락이 입력 구조에 담기지 못하는 것’에 있다.", bold_lead="AI: ")
p(doc, "→ 몰라서 생긴 차이. 채널을 바꿔도 ‘여기/아까 그거’가 무엇을 가리키는지는 그대로 미해결로 "
       "남는다는 걸 몰랐다.", space_after=6)
p(doc, "Gap 2. 해결의 방향", bold=True, space_after=3)
bullet(doc, "입력을 더 쉽게 만든다(음성).", bold_lead="나의 시도: ")
bullet(doc, "입력의 ‘구조’를 바꾼다 — 질문을 시간 위의 한 지점에 핀으로 고정한다.", bold_lead="AI: ")
p(doc, "→ 관점이 달라서 생긴 차이. 단, AI 관점이 더 설명력이 있다. 같은 답변이라도 ‘시간 맥락’에 붙으면 "
       "듣기와 질문 만들기가 충돌하지 않는다.", space_after=6)
p(doc, "Gap 3. 사용자 장면의 해상도", bold=True, space_after=3)
bullet(doc, "‘강의 들으며 질문한다’ 정도의 넓은 장면.", bold_lead="나의 시도: ")
bullet(doc, "‘질문을 만드는 순간 현재 흐름을 놓친다’는 충돌의 한 컷.", bold_lead="AI: ")
p(doc, "→ 빠뜨려서 생긴 차이. 나도 그 경험이 있었지만, 그것을 마찰의 핵심 증거로 꺼내지 못했다.", space_after=6)

section_bar(doc, "4. Takeaway: 1~3에서 배운 점")
p(doc, "이 과제를 처음부터 다시 한다면, 다음 도구들을 출발선에 두겠다.", space_after=4)
bullet(doc, "좋은 HID 문제는 “질문을 쉽게 말하게 하기”가 아니라 “질문이 발생한 맥락을 입력 구조로 "
            "만들기”다.", bold_lead="정의: ")
bullet(doc, "입력 채널만 바꾼 해결책(텍스트→음성)은 ‘인터랙션 구조 변경’이 아니다. "
            "“다른 입력 구조였다면 사라졌을 마찰인가?”를 통과해야 진짜 후보다.", bold_lead="기준: ")
bullet(doc, "마찰은 ‘입력 노동’이 아니라 ‘충돌하는 두 행위’로 잡을 때 선명해진다(듣기 vs 질문 만들기).",
       bold_lead="기준: ")
bullet(doc, "매 후보마다 “이 입력은 시간/공간/대상 중 무엇에 붙어 있는가?”를 묻는다. "
            "붙을 곳이 없으면 그것은 여전히 ‘맥락 없는 독립 텍스트’다.", bold_lead="질문: ")

doc.save("CoT_Note_질문북마크타임라인.docx")
print("saved CoT")
print("DONE")
