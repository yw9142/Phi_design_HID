from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent


FONT = "Malgun Gothic"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_FILL = "F2F4F7"
BORDER = "D9DEE7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = color


def set_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), BORDER)
        borders.append(tag)


def style_doc(doc, title, subtitle):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color, before, after in [
        ("Heading 1", 15, BLUE, 14, 7),
        ("Heading 2", 12.5, BLUE, 10, 5),
        ("Heading 3", 11, DARK_BLUE, 7, 3),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 28, 41)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(subtitle)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(9.5)
    run.font.color.rgb = MUTED


def add_para(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        rest = p.add_run(text[len(bold_prefix) :])
        rest.font.name = FONT
        rest._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    else:
        run = p.add_run(text)
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(10)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(10)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_borders(table)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_text(cell, header, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = widths[i]
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(cells[i], value)
    doc.add_paragraph()
    return table


def save(doc, name):
    out = ROOT / name
    doc.save(out)
    return out


def build_one_pager():
    doc = Document()
    style_doc(
        doc,
        "[1 Pager] 질문 북마크 타임라인 v0.2",
        "2주차 과제: 1주차 문제 정의를 HID 원칙과 Reference Book 기준으로 다듬은 버전",
    )

    doc.add_heading("문제 정의", level=1)
    doc.add_heading("1. 사용자와 페인포인트", level=2)
    add_para(
        doc,
        "실시간 온라인 강의, 회의, 웨비나처럼 재생 흐름을 스스로 멈추기 어려운 환경에서 LLM을 보조 학습 도구로 함께 쓰는 주니어 실무자 또는 학습자. 이 사용자는 익숙하지 않은 개념을 따라가면서도 현재 발화 흐름을 놓치면 다시 따라잡기 어렵다. 이해가 막히는 순간 바로 질문을 쓰면 다음 내용을 놓치고, 나중에 질문하려 하면 방금의 표현과 앞뒤 맥락을 작업 기억에서 다시 꺼내야 한다.",
    )
    add_para(
        doc,
        "2주차에서 더 좁힌 점: 사용자를 넓은 '온라인 학습자'가 아니라, 실시간 흐름 속에서 낯선 개념을 듣고 있고 되감기/정지가 사회적 또는 상황적으로 부담스러운 사람으로 한정했다. 이 조건 때문에 문제는 단순 입력 귀찮음이 아니라 시간 맥락을 잃는 문제로 생긴다.",
    )

    doc.add_heading("2. 사용자의 목표", level=2)
    add_para(
        doc,
        "실시간으로 흐르는 강의/회의의 청취 흐름을 유지하면서, 이해가 막힌 시점을 그 순간의 transcript와 함께 표시하고, 끝난 뒤 또는 잠깐 여유가 생긴 때 해당 맥락 위에서 질문해 이해를 회수한다. 완료 조건은 관찰 가능하다: 막힌 지점 표시 -> 시간 맥락 카드 생성 -> 질문 후보 선택/수정 -> 답변 카드 부착 -> 복습 타임라인에서 해결/미해결 상태 확인.",
    )

    doc.add_heading("3. 핵심 마찰 정의", level=2)
    add_para(
        doc,
        "채팅 UX는 질문을 '맥락 없는 독립 텍스트'로 받기 때문에, 질문이 발생한 시간 맥락을 입력 구조 안에 보존하지 못한다. 사용자는 질문을 쓰는 동안 직전 20~30초 맥락을 작업 기억에 붙잡아야 하고, 동시에 현재 흘러가는 내용을 계속 들어야 한다. 그 결과 듣기와 질문 만들기가 충돌한다.",
    )
    add_para(
        doc,
        "Reference Book 기준으로 보면 이 마찰은 Working Memory와 Encoding Specificity의 문제다. 사람은 한 번에 많은 정보를 오래 들고 있을 수 없고, 기억은 저장될 때의 맥락 단서와 함께 다시 제시될 때 더 잘 인출된다. 따라서 사용자가 나중에 '아까 그 부분'을 기억해 설명하게 만드는 채팅 구조는 인터페이스 문제다. 모델 답변 품질이나 음성 입력 부재가 핵심 원인이 아니다.",
    )

    doc.add_heading("해결책", level=1)
    doc.add_heading("4. 인터랙션 설계", level=2)
    add_para(
        doc,
        "질문 입력을 '문장 작성'에서 '시간 위의 한 지점 표시'로 바꾼다. 사용자는 먼저 질문을 완성하지 않고 막힌 순간만 표시한다. 시스템은 그 시점의 transcript, 개념 라벨, 맥락 요약을 함께 저장하고, 나중에 질문 후보를 제한된 수로 제시한다.",
    )
    add_table(
        doc,
        ["상태", "사용자 입력 / 조건", "시스템 출력"],
        [
            ["Listening", "강의/회의를 듣는 중", "현재 transcript와 진행 시간이 흐르고, 사용자는 막힌 순간을 기다린다."],
            ["Marked", "막힘 핀 찍기 또는 Q키", "흐름은 멈추지 않고 직전 20~30초 구간, 개념 라벨, 맥락 요약을 카드로 저장한다."],
            ["Question Candidate", "사용자가 카드 확인", "저장된 맥락을 먼저 보여주고, 그 맥락에 맞는 질문 후보 3개만 제시한다."],
            ["Answered", "질문 후보 선택 또는 한 줄 수정", "답변이 선형 채팅 로그가 아니라 해당 타임스탬프 카드 아래에 붙는다."],
            ["Review", "복습 모드 진입", "막힌 지점, 해결 여부, 질문/답변을 시간순으로 다시 본다."],
        ],
        [Inches(1.1), Inches(1.75), Inches(3.65)],
    )
    add_para(
        doc,
        "설계 원칙 반영: Hick's Law를 따라 질문 후보는 3개로 제한하고, Progressive Disclosure를 따라 직접 수정 입력은 카드 단계에서만 보인다. System Status Visibility를 위해 상단에 놓친 자막 수, 저장된 맥락 초, 해결률을 계속 보여준다.",
    )

    doc.add_heading("5. 테스트 시나리오", level=2)
    add_numbered(
        doc,
        [
            "프로토타입을 열고 더미 강의 'A/B 테스트 기초: 통계적 유의성'을 재생한다.",
            "기존 채팅 방식으로 전환해 질문을 입력한다. 입력 중 자막이 계속 흐르고 '놓친 자막 N줄'이 증가하는지 확인한다.",
            "질문 북마크 방식으로 전환해 같은 흐름에서 막힘 핀을 누른다. 자막은 계속 흐르지만 카드에는 직전 구간이 저장되는지 확인한다.",
            "카드에서 개념 라벨, 맥락 요약, 자동 저장 구간을 확인한다. 사용자가 기억해 설명하지 않아도 질문 후보가 생성되는지 확인한다.",
            "질문 후보 하나를 선택하거나 직접 한 줄로 수정해 답변을 붙인다.",
            "복습 모드를 열어 막힌 지점이 시간순으로 정리되고 해결/미해결 상태가 드러나는지 확인한다.",
        ],
    )

    doc.add_heading("6. 프로토타입 링크", level=2)
    add_para(doc, "GitHub: https://github.com/yw9142/Phi_design_HID.git")
    add_para(doc, "Live demo: https://question-bookmark-timeline.vercel.app")
    add_para(doc, "Local path: question-bookmark-timeline/")

    doc.add_heading("버전 로그", level=1)
    add_table(
        doc,
        ["버전", "변경 내용", "이유"],
        [
            ["v0.1", "막힘 버튼, 직전 transcript 저장, 질문 후보, 답변 카드, 복습 모드 구현", "전체 사이클을 먼저 작동시키기 위한 초기 프로토타입"],
            ["v0.2", "사용자 조건을 실시간 흐름/낯선 개념/되감기 어려움으로 좁힘", "문제의 원인을 입력 귀찮음이 아니라 시간 맥락 상실로 선명하게 만들기 위해"],
            ["v0.2", "카드에 개념 라벨, 맥락 요약, 저장 구간을 추가", "Encoding Specificity 원칙처럼 저장 시점의 단서를 복습 시점에도 다시 주기 위해"],
            ["v0.2", "상단 검증 지표와 3개 질문 후보 유지", "System Status Visibility와 Hick's Law를 프로토타입에서 보이게 하기 위해"],
        ],
        [Inches(0.75), Inches(3.0), Inches(2.75)],
    )

    doc.add_heading("원칙 점검", level=1)
    add_bullets(
        doc,
        [
            "인터페이스 문제가 맞다: 답변 품질이 아니라 질문이 시간/맥락에 붙지 못하는 입력 구조의 문제다.",
            "짜증 9 문제가 맞다: 실시간 흐름을 놓치면 이후 내용을 따라가기 어렵고, 나중에 질문하면 맥락 재구성 비용이 크다.",
            "사용자 목표에 솔루션을 넣지 않았다: 목표는 흐름 유지와 이해 회수이고, 북마크 타임라인은 그 목표를 위한 수단이다.",
            "모호한 단어를 줄였다: '온라인 학습자' 대신 '재생 흐름을 스스로 멈추기 어려운 실시간 청취자'로 좁혔다.",
            "필요한 것만 담았다: STT, 자동 요약, 전체 강의 검색은 제외하고 시간 핀, 맥락 카드, 질문 후보, 복습만 남겼다.",
        ],
    )
    return save(doc, "1Pager_질문북마크타임라인_v0.2.docx")


def build_cot_note():
    doc = Document()
    style_doc(
        doc,
        "CoT Note - 질문 북마크 타임라인 v0.2",
        "Abduction -> AI/Reference -> Gap Analysis -> Takeaway",
    )

    doc.add_heading("1. Abduction: 나의 답변 먼저 써보기", level=1)
    add_para(
        doc,
        "1주차 결과물의 핵심은 '질문이 발생한 시간 맥락을 보존하지 못한다'였다. 2주차를 시작할 때 내 첫 판단은, 이미 방향은 맞지만 프로토타입이 아직 '강의 북마크/요약 앱'처럼 보일 위험이 있다는 것이었다. 막힘 버튼과 질문 후보는 작동하지만, 왜 이것이 채팅 UX의 구조 문제인지 화면과 문서에서 더 분명히 증명해야 한다고 봤다.",
    )
    add_para(
        doc,
        "처음 세운 개선 가설: 사용자를 더 좁히고, 카드에 저장 시점의 단서를 더 많이 붙이면, 단순 북마크가 아니라 '맥락을 입력 구조로 바꾸는 인터페이스'라는 점이 보일 것이다. 또한 before/after를 보려면 놓친 자막 수와 저장된 맥락 같은 관찰 가능한 지표가 필요하다고 생각했다.",
    )

    doc.add_heading("2. AI / Reference: 자료와 피드백 받기", level=1)
    add_para(
        doc,
        "HID 1 Pager Guidelines에서는 사용자 조건, 관찰 가능한 목표, 하나의 핵심 마찰, 행위->시스템 반응 쌍, 테스트 시나리오를 요구한다. 1주차 문서는 방향은 좋지만 사용자 조건이 아직 넓고, 프로토타입 설명이 기능 목록처럼 읽힐 수 있었다.",
    )
    add_para(
        doc,
        "[HID] Reference Book에서 이번 문제와 직접 연결된 기준은 Working Memory, Encoding Specificity Principle, Hick's Law, Progressive Disclosure, System Status Visibility였다. 사람은 방금 들은 내용을 오래 들고 있지 못하고, 저장할 때의 맥락 단서가 다시 제시될 때 기억을 더 잘 꺼낸다. 선택지가 많으면 결정이 느려지고, 시스템이 무엇을 저장했는지 상태를 보여줘야 사용자가 결과를 해석할 수 있다.",
    )
    add_para(
        doc,
        "AI 관점의 피드백으로 정리하면, 2주차의 핵심은 새 기능을 많이 추가하는 것이 아니라 문제의 증거를 화면에 남기는 것이다. '질문을 쉽게 입력'시키는 해결책이 아니라, '질문이 생긴 시간 맥락을 사용자의 작업 기억 밖으로 꺼내 카드로 고정'하는 해결책이어야 한다.",
    )

    doc.add_heading("3. Gap Analysis: 내 답변과 비교 분석", level=1)
    add_table(
        doc,
        ["항목", "1주차/초기 생각", "2주차에서 바꾼 점", "차이의 성격"],
        [
            [
                "사용자",
                "온라인 강의, 회의, 웨비나를 듣는 학습자/실무자",
                "실시간 흐름, 낯선 개념, 되감기 어려움이라는 조건을 가진 사용자",
                "빠뜨려서 생긴 차이. 이 조건이 있어야 왜 시간 맥락 보존이 필요한지 선명해진다.",
            ],
            [
                "마찰",
                "질문이 발생한 시간 맥락을 보존하지 못함",
                "작업 기억에 직전 맥락을 붙잡은 채 현재 흐름을 들어야 하는 충돌",
                "관점의 해상도 차이. 같은 문제를 인지 부하와 기억 인출 단서로 더 설명했다.",
            ],
            [
                "해결책",
                "막힘 버튼을 누르면 질문 카드 생성",
                "시간 핀 + transcript snapshot + 개념 라벨 + 맥락 요약 + 3개 질문 후보",
                "기능 추가가 아니라 입력 구조 명확화. 질문보다 맥락이 먼저 저장된다.",
            ],
            [
                "프로토타입 검증",
                "before/after를 사용자가 체감",
                "놓친 자막 수, 저장된 맥락 초, 해결률을 상단에 표시",
                "몰라서 생긴 차이. 상태 가시성이 있어야 사용자가 효과를 해석할 수 있다.",
            ],
        ],
        [Inches(1.05), Inches(1.65), Inches(2.05), Inches(1.75)],
    )

    doc.add_heading("4. Takeaway: 다음 시도에 가져갈 도구", level=1)
    add_bullets(
        doc,
        [
            "정의: 좋은 HID 문제는 사용자가 말을 못 하는 문제가 아니라, 사용자가 이미 가진 시간/공간/대상 맥락이 인터페이스 입력으로 보존되지 않는 문제를 찾는 것이다.",
            "기준: 목표 문장에는 솔루션을 넣지 않는다. 사용자의 목표는 '흐름을 유지하며 이해를 회수한다'이고, 타임라인은 그 목표를 위한 수단이다.",
            "기준: Working Memory에 맡기는 정보를 시스템 객체로 바꿔야 한다. 이번 경우 직전 transcript와 개념 라벨이 그 객체다.",
            "기준: 질문 후보는 많을수록 좋은 것이 아니다. Hick's Law 관점에서는 3개 정도의 선택지만 먼저 보여주고, 직접 수정은 필요할 때 열어 두는 편이 낫다.",
            "질문: 이 인터페이스에서 사용자가 입력하는 것은 텍스트인가, 아니면 시간/공간/대상에 붙은 조작인가? 텍스트만 남으면 다시 채팅 UX로 돌아간 것이다.",
        ],
    )

    doc.add_heading("이번 v0.2에 실제 반영한 것", level=1)
    add_numbered(
        doc,
        [
            "1 Pager 문제 정의를 '실시간 흐름 속 낯선 개념을 듣는 사용자'로 좁혔다.",
            "핵심 마찰에 Working Memory와 Encoding Specificity 근거를 추가했다.",
            "프로토타입 카드에 개념 라벨, 맥락 요약, 자동 저장 구간을 추가했다.",
            "상단에 놓친 자막 수, 보존 맥락 초, 해결률을 보여주는 검증 지표를 추가했다.",
            "질문 후보는 3개로 유지하고, 직접 수정은 카드 내부 보조 입력으로 남겼다.",
        ],
    )
    return save(doc, "CoT_Note_질문북마크타임라인_v0.2.docx")


if __name__ == "__main__":
    print(build_one_pager())
    print(build_cot_note())
